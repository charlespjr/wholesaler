#!/usr/bin/env python3
"""
Generate a Paragon Government Solutions offer letter (Letter of Intent)
from the master template, preserving the branded header (logo) and footer
(company info, UEI, CAGE code, website).

The master template lives at templates/offer_letter_template.docx. This
script clones it, strips the body paragraphs, and rebuilds the body with
the deal-specific content while leaving the header + footer untouched.

Usage (edit the DEAL dict below, then run):
    python3 tools/make_offer_letter.py
"""
import copy
import datetime
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

TEMPLATE = "templates/offer_letter_template.docx"

# ---- Buyer constants (rarely change) ----
BUYER = {
    "entity": "Paragon Government Solutions LLC",
    "signatory": "Charles Pleasant",
    "title": "Managing Member",
    "address": "11166 Fairfax Blvd, Suite 500, Fairfax, VA 22030",
    "phone": "(888) 495-6935",
    "email": "charles@paragongovsolutions.net",
}


def _amount_words(n):
    """Quick spell-out for typical whole-dollar offers (best effort)."""
    try:
        import inflect
        p = inflect.engine()
        return p.number_to_words(int(n), andword="").title()
    except Exception:
        return ""


def num(p):
    return p.paragraph_format


def style_run(r, bold=False, size=11, font=None):
    r.bold = bold
    r.font.size = Pt(size)
    if font:
        r.font.name = font
    return r


def add_hrule(doc):
    p = doc.add_paragraph()
    num(p).space_after = Pt(12)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "2E4057")
    pBdr.append(bottom)
    pPr.append(pBdr)


def labeled(doc, label, value):
    p = doc.add_paragraph()
    num(p).space_after = Pt(4)
    num(p).left_indent = Inches(0.3)
    style_run(p.add_run(f"{label}  "), bold=True)
    style_run(p.add_run(value))


def section_head(doc, text):
    p = doc.add_paragraph()
    num(p).space_after = Pt(4)
    style_run(p.add_run(text), bold=True)


def para(doc, text="", bold=False, size=11, align=WD_ALIGN_PARAGRAPH.LEFT,
         space_after=6, font=None):
    p = doc.add_paragraph()
    p.alignment = align
    num(p).space_after = Pt(space_after)
    num(p).space_before = Pt(0)
    if text:
        style_run(p.add_run(text), bold=bold, size=size, font=font)
    return p


def build(deal, out_path):
    doc = Document(TEMPLATE)

    # Wipe existing body paragraphs (header/footer are separate, stay intact)
    for p in list(doc.paragraphs):
        p._element.getparent().remove(p._element)

    today = deal.get("date") or datetime.date.today().strftime("%B %-d, %Y")

    # --- Header block (body text, mirrors template) ---
    para(doc, BUYER["entity"].upper(), bold=True, size=14,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    para(doc, "11166 Fairfax Blvd, Suite 500  |  Fairfax, VA 22030", size=10,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    para(doc, f"Phone: {BUYER['phone']}  |  Email: {BUYER['email']}", size=10,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=16)
    add_hrule(doc)

    para(doc, f"Date: {today}", space_after=12)
    para(doc, "LETTER OF INTENT TO PURCHASE REAL PROPERTY", bold=True, size=13,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=16)

    para(doc,
         f'{BUYER["entity"]} ("Buyer") hereby submits this Letter of Intent '
         "to purchase the real property described below. This letter outlines "
         "our proposed terms and is intended to serve as the basis for a "
         "formal Purchase and Sale Agreement.", space_after=10)
    add_hrule(doc)

    # --- 1. Property ---
    section_head(doc, "1.  PROPERTY")
    labeled(doc, "Property Address:", deal["property_address"])
    labeled(doc, "Legal Description:",
            deal.get("legal", "As described in public records — Parcel / "
                              "MLS listing per listing agent."))
    labeled(doc, "Property Type:", deal.get("property_type",
                                            "Residential Single Family"))
    labeled(doc, "Current Listing Agent:", deal["agent"])
    para(doc, space_after=6)

    # --- 2. Buyer ---
    section_head(doc, "2.  BUYER")
    labeled(doc, "Entity Name:", BUYER["entity"])
    labeled(doc, "Authorized Signatory:", BUYER["signatory"])
    labeled(doc, "Address:", BUYER["address"])
    labeled(doc, "Phone:", BUYER["phone"])
    labeled(doc, "Email:", BUYER["email"])
    para(doc, space_after=6)

    # --- 3. Price & Terms ---
    section_head(doc, "3.  PURCHASE PRICE & TERMS")
    words = _amount_words(deal["price"])
    price_str = f"${deal['price']:,}"
    if words:
        price_str += f" ({words} Dollars)"
    labeled(doc, "Purchase Price:", price_str)
    labeled(doc, "Financing:", deal.get("financing",
                                        "ALL CASH — No financing contingency"))
    labeled(doc, "Earnest Money Deposit:", deal.get("emd",
            "$1,000 deposited within 3 business days of fully executed contract"))
    labeled(doc, "Property Condition:", deal.get("condition",
            "As-Is — Buyer waives all inspection and repair contingencies."))
    labeled(doc, "Closing Costs:", deal.get("closing_costs",
            f"Split per customary {deal.get('state','local')} practice"))
    labeled(doc, "Closing Timeline:", deal.get("closing_timeline",
            "30 days after fully executed Purchase & Sale Agreement"))
    labeled(doc, "Contingencies:", deal.get("contingencies",
            "None beyond clear, marketable title"))
    labeled(doc, "Proof of Funds:", deal.get("pof", "Available upon request"))
    para(doc, space_after=6)

    # --- 4. Additional Terms ---
    section_head(doc, "4.  ADDITIONAL TERMS")
    for i, t in enumerate(deal["additional_terms"], 1):
        p = doc.add_paragraph()
        num(p).space_after = Pt(6)
        num(p).left_indent = Inches(0.3)
        style_run(p.add_run(f"{i}.  {t}"))
    para(doc, space_after=10)
    add_hrule(doc)

    para(doc, "We appreciate the opportunity to submit this offer and look "
              "forward to working with you toward a smooth and timely closing. "
              "Please do not hesitate to contact us with any questions.",
         space_after=12)

    # --- Signature ---
    para(doc, "Sincerely,", space_after=24)
    para(doc, BUYER["signatory"], size=20, font="Blank Script", space_after=2)
    para(doc, BUYER["signatory"], bold=True, space_after=2)
    para(doc, BUYER["title"], space_after=2)
    para(doc, BUYER["entity"], space_after=2)
    para(doc, f"{BUYER['phone']}  |  {BUYER['email']}", size=10, space_after=0)

    doc.save(out_path)
    print("Saved:", out_path)


# ---------------------------------------------------------------------------
# Edit this block per deal, then run the script.
# ---------------------------------------------------------------------------
DEAL = {
    "property_address": "6006 Groom Rd, Baker, LA 70714",
    "agent": "Laurie Rush | laurie@kaizenhomesales.com",
    "price": 42500,
    "state": "Louisiana",
    "closing_timeline": "30 days after bank approval and fully executed "
                        "Purchase & Sale Agreement",
    "closing_costs": "Split per customary Louisiana practice",
    "additional_terms": [
        "Buyer understands this is a bank-owned (REO) property and acknowledges "
        "that additional time may be required for the selling institution to "
        "review and approve this offer.",
        "Buyer is prepared to sign the seller's standard REO Purchase and Sale "
        "Agreement and addenda upon acceptance.",
        "This Letter of Intent is non-binding and is subject to the execution "
        "of a formal Purchase and Sale Agreement signed by all parties.",
        "This offer is valid for 14 days from the date above.",
    ],
}

if __name__ == "__main__":
    build(DEAL, "offer_letter_6006_Groom_Rd.docx")
